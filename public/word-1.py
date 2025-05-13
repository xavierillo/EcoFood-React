from docx import Document
from docx.shared import Pt, Inches

# Create document
doc = Document()

# Title
doc.add_heading('EcoFood – Guía de Primeros Pasos\nUnidad 3 · Framework basado en JavaScript', 0)

# Introducción
doc.add_heading('1. Introducción y objetivo', level=1)
p = doc.add_paragraph(
    'Este documento resume los pasos iniciales que cada equipo (3‑4 estudiantes) '
    'debe seguir durante la semana 1 para levantar la base técnica de la aplicación '
    'EcoFood con React + Firebase. Al finalizar estos pasos tendrás un proyecto '
    'ejecutable, versionado en Git y enlazado a un backend gratuito en Firebase.')

# Arquitectura
doc.add_heading('2. Patrón de arquitectura recomendado', level=1)
doc.add_paragraph(
    'En lugar del patrón MVC tradicional, usaremos un enfoque **feature‑based / '
    'component‑driven**, donde la “vista” y parte de la lógica conviven dentro de '
    'componentes React. Las funcionalidades se agrupan por dominio, lo que facilita '
    'escalar, testear y repartir el trabajo entre integrantes.')
doc.add_paragraph('Estructura sugerida:', style='Intense Quote')
structure = (
    'src/\n'
    '├── api/            # Peticiones y helpers de red\n'
    '├── assets/         # Imágenes globales y fuentes\n'
    '├── components/     # Piezas reutilizables (Navbar, CardProducto…)\n'
    '├── hooks/          # useAuth, useProductos, …\n'
    '├── context/        # Providers de estado global (AuthProvider…)\n'
    '├── pages/          # Vistas mapeadas a rutas (Login.jsx, Home.jsx…)\n'
    '├── routes/         # Definición de rutas y ProtectedRoute.jsx\n'
    '├── services/       # firebase.js y otros SDKs\n'
    '└── main.jsx        # Punto de entrada'
)
doc.add_paragraph(structure)

doc.add_paragraph(
    'Ventajas frente a MVC:\n'
    '• Escalabilidad: cada feature crece aislada.\n'
    '• Testing sencillo: se pueden «mockear» módulos dentro de /api o /services.\n'
    '• Claridad de responsabilidades: UI, estado y red están separados, pero cercanos '
    'a la funcionalidad.')

# Paso a paso
doc.add_heading('3. Pasos iniciales', level=1)

# 3.1 proyecto vite
doc.add_heading('Paso 1 – Crear proyecto React con Vite', level=2)
doc.add_paragraph('Ejecuta en terminal:', style='List Number')
doc.add_paragraph('npm create vite@latest ecofood-react -- --template react', style='List Continue')
doc.add_paragraph('cd ecofood-react', style='List Continue')
doc.add_paragraph('npm install', style='List Continue')
doc.add_paragraph('npm run dev  # abre http://localhost:5173', style='List Continue')

# 3.2 dependencias
doc.add_heading('Paso 2 – Añadir dependencias y herramientas de calidad', level=2)
deps = (
    'npm i react-router-dom\n'
    'npm i firebase\n'
    'npm i --save-dev eslint prettier eslint-config-prettier\n'
    'npm i --save-dev @testing-library/react jest'
)
doc.add_paragraph('Comandos sugeridos:', style='Intense Quote')
doc.add_paragraph(deps)

# 3.3 estructura carpetas
doc.add_heading('Paso 3 – Crear carpetas base', level=2)
doc.add_paragraph(
    'Crea manualmente las carpetas indicadas en la sección de arquitectura y mueve '
    '`App.jsx` a `src/pages/Home.jsx`.')

# 3.4 router
doc.add_heading('Paso 4 – Configurar React Router DOM', level=2)
doc.add_paragraph(
    'En `src/routes/router.jsx` declara las rutas `/`, `/login` y `/register`. Exporta '
    'un componente `<BrowserRouter>` que envuelva tu App.')

# 3.5 firebase cuenta
doc.add_heading('Paso 5 – Crear cuenta y proyecto Firebase', level=2)
steps_firebase = [
    'Inicia sesión con tu cuenta Google en https://console.firebase.google.com.',
    'Haz clic en **“Agregar proyecto”** → nómbralo *EcoFood*.',
    'Desactiva Google Analytics (opcional) y crea el proyecto.',
    'En el panel principal, selecciona el ícono `</>` para registrar la **app web** '
    '(*ecofood-web*).',
    'Copia el snippet de configuración (apiKey, authDomain, etc.).'
]
for s in steps_firebase:
    doc.add_paragraph(s, style='List Bullet')

# 3.6 firebase sdk
doc.add_heading('Paso 6 – Inicializar SDK en `firebase.js`', level=2)
code_block = (
    'import { initializeApp } from "firebase/app";\n'
    'import { getAuth } from "firebase/auth";\n'
    'import { getFirestore } from "firebase/firestore";\n\n'
    'const firebaseConfig = {\n'
    '  apiKey: import.meta.env.VITE_FIREBASE_API_KEY,\n'
    '  authDomain: import.meta.env.VITE_FIREBASE_AUTH_DOMAIN,\n'
    '  projectId: import.meta.env.VITE_FIREBASE_PROJECT_ID,\n'
    '  storageBucket: import.meta.env.VITE_FIREBASE_STORAGE_BUCKET,\n'
    '  messagingSenderId: import.meta.env.VITE_FIREBASE_MESSAGING_SENDER_ID,\n'
    '  appId: import.meta.env.VITE_FIREBASE_APP_ID,\n'
    '};\n\n'
    'const app = initializeApp(firebaseConfig);\n'
    'export const auth = getAuth(app);\n'
    'export const db = getFirestore(app);'
)
doc.add_paragraph('Ejemplo:', style='Intense Quote')
doc.add_paragraph(code_block)

# Checklist
doc.add_heading('4. Checklist de entrega (Semana 1)', level=1)
check_items = [
    'Proyecto arranca en http://localhost:5173',
    'Repositorio subido (GitHub o GitLab) con primer commit',
    'Estructura de carpetas y rutas básicas creadas',
    '.env configurado (no subido) y .env.example en el repo',
    'Proyecto Firebase enlazado a la app (firebase.js sin errores)'
]
for item in check_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph('\n¡Éxitos y a codificar!')

# Save document
file_path = '/mnt/data/Guia_Inicial_EcoFood_React_Firebase.docx'
doc.save(file_path)

file_path
